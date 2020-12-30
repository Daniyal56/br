from docx.shared import Mm,Inches,Pt
import docx
from docx2pdf import convert
import pythoncom
import csv,win32com.client,datetime,os

# wrd = win32com.client.gencache.EnsureDispatch("Word.Application")



def file_converter(csv_file):
    with open(csv_file, newline='') as f:
        csv_reader = csv.reader(f)
        # ---------------------------------------------------------------------------------------------------------------
        doc = docx.Document()

        # doc.add_paragraph(f'Date: {datetime.datetime.now()}')
        section = doc.sections[0]
        section.page_height = Mm(350)
        section.page_width = Mm(297)
        header = section.header
        heading = header.paragraphs[0]
        heading.text = f'BANK RECONCILIATION'
        date_time = header.add_paragraph()
        date_time.text = f'\n\nReconcile Date: {datetime.datetime.now()}\n\n'
        heading.style.font.size = Pt(30)
        heading.style.font.underline = True

        heading.alignment = 1
        footer = section.footer
        # footer_1 = section.footer
        foot = footer.paragraphs[0]
        f = footer.add_paragraph()
        f.text = 'Printed By SC\t\t\t\t\t\t\t\t\tThis is a system generated voucher'
        f.alignment = 1
        # -----------------------------------------------------------------------------------------------------
        header = next(csv_reader)
        cols = len(header)
        table = doc.add_table(rows=1, cols=cols, style='Table Grid')
        hdr_cells = table.rows[0].cells

        for i in range(cols):
            hdr_cells[i].text = header[i]

        for row in csv_reader:
            row_cells = table.add_row().cells
            for i in range(cols):
                row_cells[i].text = row[i]
        for cell in table.rows[0].cells:
            cell.width = Inches(0.5)
        pythoncom.CoInitialize()
        doc.save(os.path.join(f"C:/Users/Daniyal/Desktop/Bank Reconcilation/docx_folder/{csv_file}.docx"))
        convert(os.path.join(f"C:/Users/Daniyal/Desktop/Bank Reconcilation/docx_folder/{csv_file}.docx"), os.path.join(f"C:/Users/Daniyal/Desktop/Bank Reconcilation/pdf_folder/{csv_file}_.pdf"))
