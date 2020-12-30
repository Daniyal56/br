import docx
from docx.shared import Mm,Inches,Pt
from docx2pdf import convert
import pythoncom
import csv,win32com.client,datetime

wrd = win32com.client.gencache.EnsureDispatch("Word.Application")

doc= docx.Document()
# heading = doc.add_heading('Bank Reconcilation')
# heading.style.font.size=Pt(30)
# # print(dir(heading.style))
# heading.style.font.underline = True
# # p.add_run('italic.').italic = True
# heading.alignment = 1

doc.add_paragraph('')
doc.add_paragraph(f'Date: {datetime.datetime.now()}')
doc.add_paragraph('')
section = doc.sections[0]
section.page_height = Mm(594)
section.page_width = Mm(297)
header = section.header
heading = header.paragraphs[0]
heading.text = 'Bank Reconcilation'
heading.style.font.size=Pt(30)
heading.style.font.underline = True
heading.alignment = 1

footer = section.footer
# footer_1 = section.footer
foot = footer.paragraphs[0]
f = footer.add_paragraph()
f.text = 'Printed By Software Channel Pvt. Ltd.\t\t\t\t\t\t\t\t\tThis is a system generated voucher'
# f = footer_1.add_paragraph()
# foot.text = 'Printed By Software Channel Pvt. Ltd.'
# foot.style.font.size=Pt(7.5)
# f1 = footer.add_paragraph()
# f1.text = 'Printed By Software Channel Pvt. Ltd.'
f.alignment = 1


with open('Unmatched_Bank.csv',newline='') as f:
    csv_reader = csv.reader(f)
    header = next(csv_reader)
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols,style='Table Grid')
    hdr_cells = table.rows[0].cells

    for i in range(cols):
        hdr_cells[i].text = header[i]

    for row in csv_reader:
        row_cells = table.add_row().cells
        for i in range(cols):
            row_cells[i].text = row[i]
    for cell in table.rows[0].cells:
        cell.width = Inches(0.5)


doc.save("data.docx")
convert("data.docx","C:/Users/Daniyal/Desktop/Bank Reconcilation/backup/data.pdf")
# print(dir(pythoncom))
# print(dir(doc))
# print(dir(doc.core_properties))
